# -*- coding: utf-8 -*-
"""
Word生成模块 —— 对应 PRD 4.5
按 PRD 里给的文档结构生成：封面 / 证明材料 / 发票 / 垫付事由 / 签领表 / 活动策划
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def _kv_line(doc, label, value):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}：")
    run.bold = True
    p.add_run(str(value) if value not in (None, "") else "________")
    return p


def generate_docx(structured_data: dict, output_path: str):
    basic = structured_data["basic_info"]
    items = structured_data["items"]
    total = structured_data["total_amount"]
    purpose = structured_data.get("purpose", "")
    ocr_results = structured_data.get("ocr_results", [])
    today = date.today()

    doc = Document()
    doc.styles["Normal"].font.name = "微软雅黑"
    doc.styles["Normal"].font.size = Pt(11)

    # ---------- 1. 封面 ----------
    title = doc.add_heading("报销说明", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _kv_line(doc, "报销事由", basic.get("reimburse_reason"))
    method = basic.get("reimburse_method", "")
    payer = basic.get("payer_name", "")
    p = doc.add_paragraph()
    p.add_run("报销方式：").bold = True
    p.add_run(f"{method}给{payer}，报销总金额为{total:.2f}元")
    _kv_line(doc, "报销人姓名", payer)
    _kv_line(doc, "报销人学号", basic.get("student_id"))
    _kv_line(doc, "联系方式", basic.get("contact"))
    _kv_line(doc, "报销金额", f"{total:.2f} 元")
    _kv_line(doc, "活动时间", basic.get("activity_time"))

    doc.add_page_break()

    # ---------- 2. 证明材料 ----------
    _add_heading(doc, "二、证明材料", level=1)
    _kv_line(doc, "报销人姓名", payer)
    _kv_line(doc, "报销人学号", basic.get("student_id"))
    _kv_line(doc, "报销金额", f"{total:.2f} 元")

    doc.add_paragraph("明细：").runs[0].bold = True
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, text in enumerate(["序号", "名称明细", "单价", "数量"]):
        hdr[i].text = text

    for idx, item in enumerate(items, start=1):
        row = table.add_row().cells
        row[0].text = str(idx).zfill(2)
        row[1].text = item["name"]
        row[2].text = f"¥{item['unit_price']:.2f}"
        row[3].text = str(item["quantity"])

    _kv_line(doc, "合计", f"{total:.2f} 元")
    _kv_line(doc, "事由", basic.get("reimburse_reason"))
    _kv_line(doc, "用途", purpose)

    doc.add_paragraph("\n签名：________________")
    doc.add_paragraph(f"日期：{today.year}年{today.month}月{today.day}日")

    doc.add_page_break()

    # ---------- 3. 发票 / 购买详情 / 支付记录 ----------
    _add_heading(doc, "三、发票及凭证明细", level=1)
    invoice_idx, other_idx = 0, 0
    for r in ocr_results:
        ftype = r.get("file_type")
        if ftype == "invoice":
            invoice_idx += 1
            _add_heading(doc, f"发票 {invoice_idx}：", level=2)
        elif ftype == "payment":
            other_idx += 1
            _add_heading(doc, f"支付记录 {other_idx}：", level=2)
        else:
            other_idx += 1
            _add_heading(doc, f"订单截图 {other_idx}：", level=2)

        _kv_line(doc, "商户/渠道", r.get("merchant"))
        if r.get("amount") is not None:
            _kv_line(doc, "金额", f"¥{r['amount']:.2f}")
        _kv_line(doc, "日期", r.get("date"))
        if r.get("items"):
            for it in r["items"]:
                doc.add_paragraph(
                    f"· {it.get('name', '')} × {it.get('quantity', '')}"
                    f"（单价 ¥{it.get('unit_price', 0):.2f}）",
                    style="List Bullet",
                )

    doc.add_page_break()

    # ---------- 4. 垫付事由 ----------
    _add_heading(doc, "四、垫付事由说明", level=1)
    item_lines = "；".join(
        f"{i + 1}.购买{it['name']}{it['quantity']}份/件，花费{it['unit_price'] * it['quantity']:.2f}元"
        for i, it in enumerate(items)
    ) or "（无明细）"

    doc.add_paragraph(
        f"{basic.get('reimburse_reason', '')}报销垫付事由\n"
        f"本人{payer}（学号：{basic.get('student_id', '')}）在"
        f"{basic.get('reimburse_reason', '')}中垫付共计{total:.2f}元，具体用于以下事项：\n"
        f"{item_lines}。"
    )
    doc.add_paragraph("\n签名（垫付人本人）：________________")
    doc.add_paragraph(f"日期：{today.year}年{today.month}月{today.day}日")

    doc.add_page_break()

    # ---------- 5. 签领表 ----------
    _add_heading(doc, "五、签领表", level=1)
    sign_table = doc.add_table(rows=1, cols=4)
    sign_table.style = "Light Grid Accent 1"
    for i, text in enumerate(["姓名", "学号", "签领金额", "签名"]):
        sign_table.rows[0].cells[i].text = text
    row = sign_table.add_row().cells
    row[0].text = payer
    row[1].text = basic.get("student_id", "")
    row[2].text = f"¥{total:.2f}"
    row[3].text = ""

    doc.add_page_break()

    # ---------- 6. 活动策划 ----------
    _add_heading(doc, "六、活动策划与购买说明", level=1)
    doc.add_paragraph(purpose or "（未填写）")

    doc.save(output_path)
    return output_path
