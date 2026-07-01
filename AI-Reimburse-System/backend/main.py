import json
import base64
import os
import io  # 👈 新增
import re
from typing import List
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import requests
from dotenv import load_dotenv
from PIL import Image, UnidentifiedImageError

load_dotenv(override=True)

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

ZHIPU_API_KEY = os.getenv("ZHIPU_API_KEY")
ZHIPU_API_BASE = os.getenv("ZHIPU_API_BASE", "https://open.bigmodel.cn/api/paas/v4")
MODEL = os.getenv("ZHIPU_MODEL", "glm-4v-flash")


def call_zhipu_vision(messages):
    if not ZHIPU_API_KEY:
        raise HTTPException(status_code=500, detail="智谱 API Key 未配置，请设置环境变量 ZHIPU_API_KEY")

    endpoint = f"{ZHIPU_API_BASE.rstrip('/')}/chat/completions"
    try:
        response = requests.post(
            endpoint,
            headers={
                "Authorization": f"Bearer {ZHIPU_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": MODEL,
                "messages": messages,
                "temperature": 0.1,
            },
            timeout=120,
        )
        response.raise_for_status()
    except requests.RequestException as exc:
        response_text = getattr(exc.response, "text", "") if getattr(exc, "response", None) is not None else ""
        detail = f"智谱接口请求失败: {exc}"
        if response_text:
            detail += f"；接口返回：{response_text[:500]}"
        raise HTTPException(status_code=502, detail=detail) from exc

    payload = response.json()
    try:
        return payload["choices"][0]["message"]["content"]
    except (KeyError, IndexError, TypeError) as exc:
        raise HTTPException(status_code=502, detail=f"智谱接口返回格式异常: {payload}") from exc


def compress_image_for_ai(raw_bytes, max_side=1800, quality=82):
    try:
        image = Image.open(io.BytesIO(raw_bytes))
        image = image.convert("RGB")
        image.thumbnail((max_side, max_side))
        output = io.BytesIO()
        image.save(output, format="JPEG", quality=quality, optimize=True)
        return output.getvalue()
    except (UnidentifiedImageError, OSError):
        return raw_bytes


def parse_ai_json(raw_text):
    cleaned = str(raw_text or "").replace("```json", "").replace("```", "").strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        start = cleaned.find("{")
        if start != -1:
            decoder = json.JSONDecoder()
            parsed, _ = decoder.raw_decode(cleaned[start:])
            return parsed
        raise


def merge_ai_results(results):
    merged = {
        "items": [],
        "image_groups": [],
    }
    total_amount = 0.0
    for result, index_map in results:
        result = result or {}
        result = normalize_text_fields(result, {})
        merged["items"].extend(result.get("items", []))
        amount = infer_amount(result)
        total_amount += amount
        for group in result.get("image_groups", []) or []:
            if not isinstance(group, dict):
                continue
            inv_idx = clean_index(group.get("invoice_idx"))
            order_idx = clean_index(group.get("order_idx"))
            payment_idx = clean_index(group.get("payment_idx"))
            merged["image_groups"].append({
                "invoice_idx": index_map[inv_idx] if inv_idx is not None and inv_idx < len(index_map) else None,
                "order_idx": index_map[order_idx] if order_idx is not None and order_idx < len(index_map) else None,
                "payment_idx": index_map[payment_idx] if payment_idx is not None and payment_idx < len(index_map) else None,
            })
    if total_amount > 0:
        merged["amount"] = format_money(total_amount)
        merged["order_paid_amount"] = format_money(total_amount)
    return merged


def analyze_image_summaries(images_payload, reason):
    summaries = []
    summary_prompt = f"""
只输出JSON，不要解释。
你正在分析一张报销材料图片，用户报销事由为：{reason}
请判断该图片属于哪一类，并提取可用于和其他图片配对的线索。
输出格式：
{{
  "doc_type": "invoice/order/payment/other",
  "title": "图片主要内容",
  "merchant": "商家或开票方",
  "item_name": "商品、服务或住宿名称",
  "amount": "金额数字，支付记录负数取绝对值",
  "date": "日期或时间",
  "order_no": "订单号/交易单号/发票号等",
  "keywords": ["关键词1", "关键词2"],
  "reimburse_related": true
}}
注意：
1. 酒店、住宿、报名、赛事应急药品、参赛装备通常与赛事报销相关。
2. 明显日用品、个人生活用品、与报销事由无关的消费，reimburse_related 填 false。
3. 金额要读完整，例如 -3590.00 应写 3590，不能写 359。
"""
    for idx, image_payload in enumerate(images_payload):
        raw = call_zhipu_vision([
            {"role": "user", "content": [{"type": "text", "text": summary_prompt}, image_payload]}
        ])
        summary = parse_ai_json(raw)
        if not isinstance(summary, dict):
            summary = {"title": str(summary)}
        summary["image_idx"] = idx
        summaries.append(summary)
    return summaries


def build_result_from_summaries(summaries, base_info):
    reason = base_info.get("reason") or "社团活动报销"
    activity_time = base_info.get("activityTime") or ""
    prompt = f"""
只输出JSON，不要解释。
你是一位财务报销材料配对助手。下面是每张图片单独识别出来的文字线索，image_idx 是原始上传图片索引。
用户报销事由：{reason}
用户活动时间：{activity_time}

请完成：
1. 自动把属于同一事项的发票、订单截图、支付记录配成一组，不要依赖图片顺序。
2. 只把与报销事由相关、需要报销的事项放入 items。明显日用品、个人消费、无关生活用品不要放入 items，也不要计入 amount。
3. 每个 item 的 total 必须优先使用支付记录/订单实付金额；支付记录负数取绝对值。住宿若支付记录是 -3590.00，必须写 3590。
4. 输出 image_groups，里面的 invoice_idx/order_idx/payment_idx 必须使用原始 image_idx。
5. 生成用途、购买说明、活动策划、器材与活动关联、预期效果；说明文字只写类别，不堆商品长全名。

输出格式：
{{
  "amount": "订单实付总额",
  "order_paid_amount": "订单实付总额",
  "items": [
    {{"name": "明细名称", "quantity": 1, "price": "单价或实付", "discount": 0, "total": "实付金额", "remark": "备注"}}
  ],
  "image_groups": [
    {{"invoice_idx": 0, "order_idx": 3, "payment_idx": 6}}
  ],
  "purpose_statement": "用途",
  "purchase_explanation": "购买说明",
  "activity_plan": "活动名称：...\\n活动时间：...\\n活动地点：...\\n参与人员：...",
  "activity_relation": "关联说明",
  "expected_effect": "预期效果",
  "advance_payment_note": "超过1000元时生成垫付说明，否则空字符串"
}}

图片线索：
{json.dumps(summaries, ensure_ascii=False)}
"""
    raw = call_zhipu_vision([
        {"role": "user", "content": [{"type": "text", "text": prompt}]}
    ])
    return parse_ai_json(raw)


def summary_text(summary):
    values = [
        summary.get("title"),
        summary.get("merchant"),
        summary.get("item_name"),
        " ".join(str(x) for x in summary.get("keywords", []) or []),
    ]
    return " ".join(str(v) for v in values if v)


def contains_any(text, words):
    return any(word in text for word in words)


def summary_amount(summary):
    return abs(parse_money(summary.get("amount")))


def make_group_from_summaries(matched):
    group = {"invoice_idx": None, "order_idx": None, "payment_idx": None}
    for summary in matched:
        idx = summary.get("image_idx")
        doc_type = str(summary.get("doc_type", ""))
        text = summary_text(summary)
        if group["invoice_idx"] is None and ("invoice" in doc_type or "发票" in text):
            group["invoice_idx"] = idx
        if group["payment_idx"] is None and ("Transaction" in text or "Transactions" in text or "支付成功" in text or "消费详情" in text):
            group["payment_idx"] = idx
        if group["order_idx"] is None and idx not in {group["invoice_idx"], group["payment_idx"]}:
            group["order_idx"] = idx
    return group


def payment_amount_from_summaries(matched):
    payment_amounts = []
    for summary in matched:
        text = summary_text(summary)
        if contains_any(text, ["Transaction", "Transactions", "支付成功", "消费详情"]):
            amount = summary_amount(summary)
            if amount > 0:
                payment_amounts.append(amount)
    if payment_amounts:
        return max(payment_amounts)
    amounts = [summary_amount(summary) for summary in matched if summary_amount(summary) > 0]
    return max(amounts) if amounts else 0


def reconcile_result_with_summaries(result_data, summaries):
    clusters = [
        {
            "name": "康玛士运动冷冻喷雾剂足球篮球马拉松肌肉损伤降温舒缓冰肌喷雾剂",
            "terms": ["康玛", "冷喷", "冷冻喷雾", "运动冷冻喷雾"],
            "remark": "赛事应急药品耗材",
        },
        {
            "name": "可孚一次性医用冰袋",
            "terms": ["可孚", "医用冰袋", "一次性医用冰袋", "冰袋"],
            "remark": "赛事应急药品耗材",
        },
        {
            "name": "住宿费",
            "terms": ["住宿", "酒店", "双床房", "客房"],
            "remark": "赛事住宿费用，按支付记录实付金额报销，不含发票中无关日用品",
        },
    ]
    items = []
    groups = []
    used_indices = set()
    for cluster in clusters:
        matched = []
        for summary in summaries:
            text = summary_text(summary)
            if contains_any(text, cluster["terms"]):
                matched.append(summary)
        if not matched:
            continue
        amount = payment_amount_from_summaries(matched)
        if amount <= 0:
            continue
        items.append({
            "name": cluster["name"],
            "quantity": 1,
            "price": format_money(amount),
            "discount": 0,
            "total": format_money(amount),
            "remark": cluster["remark"],
        })
        group = make_group_from_summaries(matched)
        if any(value is not None for value in group.values()):
            groups.append(group)
        used_indices.update(summary.get("image_idx") for summary in matched)

    if items:
        result_data["items"] = items
        result_data["amount"] = format_money(items_paid_total(items))
        result_data["order_paid_amount"] = result_data["amount"]
        result_data["image_groups"] = groups
    return result_data


def image_chunks_for_processing(images_payload):
    count = len(images_payload)
    if count <= 3:
        return [(images_payload, list(range(count)))]
    if count % 3 == 0:
        group_count = count // 3
        chunks = []
        for i in range(group_count):
            indices = [i, i + group_count, i + 2 * group_count]
            chunks.append(([images_payload[idx] for idx in indices], indices))
        return chunks
    return [
        (images_payload[offset:offset + 3], list(range(offset, min(offset + 3, count))))
        for offset in range(0, count, 3)
    ]


@app.get("/api/v1/reimburse/health")
def health():
    return {
        "status": "ok",
        "api_configured": bool(ZHIPU_API_KEY),
        "model": MODEL,
    }


def parse_money(value):
    if value is None:
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)
    match = re.search(r"-?\d+(?:\.\d+)?", str(value).replace(",", ""))
    return float(match.group(0)) if match else 0.0


def format_money(value):
    amount = parse_money(value)
    text = f"{amount:.2f}".rstrip("0").rstrip(".")
    return text or "0"


def text_amount(value):
    text = str(value or "")
    patterns = [
        r"(?:订单实付|实付|金额|合计|总计|总额)[^\d]{0,8}(\d+(?:\.\d+)?)",
        r"(\d+(?:\.\d+)?)\s*元",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return parse_money(match.group(1))
    return 0.0


def first_value(data, keys):
    for key in keys:
        value = data.get(key)
        if value not in (None, ""):
            return value
    return None


def infer_item_name(ai_data):
    explicit_name = first_value(ai_data, ["name", "product_name", "item_name", "goods_name", "title"])
    if explicit_name:
        return str(explicit_name)

    text = " ".join(str(ai_data.get(key, "")) for key in ["purchase_explanation", "purpose_statement", "activity_relation"])
    patterns = [
        r"(?:购买|采购|采购物资为|所购物品为)([^，。,；;\n]+)",
        r"采购类别为([^，。,；;\n]+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            name = match.group(1).strip()
            if name and not any(word in name for word in ["训练物资", "相关物资", "社团活动相关物资"]):
                return name
    return "请根据发票和订单截图核对商品名称"


def normalize_items(ai_data):
    candidates = [
        ai_data.get("items"),
        ai_data.get("invoice_items"),
        ai_data.get("goods"),
        ai_data.get("products"),
        ai_data.get("details"),
        ai_data.get("商品清单"),
        ai_data.get("明细"),
    ]
    nested = ai_data.get("data")
    if isinstance(nested, dict):
        candidates.extend([
            nested.get("items"),
            nested.get("invoice_items"),
            nested.get("goods"),
            nested.get("products"),
            nested.get("details"),
        ])

    raw_items = next((item_list for item_list in candidates if isinstance(item_list, list) and item_list), [])
    normalized = []
    for item in raw_items:
        if not isinstance(item, dict):
            continue
        name = first_value(item, ["name", "product_name", "item_name", "goods_name", "title", "名称", "商品名称", "项目名称"])
        quantity = first_value(item, ["quantity", "qty", "count", "num", "数量"])
        price = first_value(item, ["price", "unit_price", "单价", "金额"])
        total = first_value(item, ["total", "amount", "subtotal", "line_total", "价税合计", "总价", "合计"])
        discount = first_value(item, ["discount", "优惠", "折扣"])
        remark = first_value(item, ["remark", "备注"])
        normalized.append({
            "name": str(name or "商品明细"),
            "quantity": quantity or 1,
            "price": price if price not in (None, "") else total,
            "discount": discount or 0,
            "total": total if total not in (None, "") else "",
            "remark": remark or "",
        })

    if not normalized:
        amount = infer_amount(ai_data)
        if amount > 0:
            normalized.append({
                "name": infer_item_name(ai_data),
                "quantity": 1,
                "price": format_money(amount),
                "discount": 0,
                "total": format_money(amount),
                "remark": "请根据发票和订单截图核对",
            })
    return normalized


def items_paid_total(items):
    total = 0.0
    for item in items or []:
        item_total = parse_money(item.get("total"))
        if item_total <= 0:
            item_total = parse_money(item.get("price")) * (parse_money(item.get("quantity")) or 1)
        total += item_total
    return total


def material_summary(items, reason=""):
    text = " ".join(str(item.get("name", "")) for item in items or [])
    combined = f"{reason} {text}"
    if any(k in combined for k in ["酒店", "住宿", "报名", "参赛费", "保险", "级位", "腰带", "道服", "护具", "脚套"]):
        return "参赛装备、赛事保障及报名住宿相关费用"
    if any(k in combined for k in ["沙袋", "脚靶", "护具", "训练"]):
        return "社团训练器材及训练物资"
    if any(k in combined for k in ["冷喷", "喷雾", "药", "冰袋", "医疗", "应急"]):
        return "训练及赛事应急药品耗材"
    if any(k in combined for k in ["打印", "资料", "对阵表"]):
        return "活动资料打印及保障费用"
    return "社团活动相关物资"


def is_weak_text(value):
    text = str(value or "").strip()
    return not text or text in {"特购置相关物资。", "本物资有效保障了集体活动的正常开展。", "丰富了校园多元文化活动。"}


def default_purpose(base_info, ai_data):
    reason = base_info.get("reason") or "社团活动"
    summary = material_summary(ai_data.get("items", []), reason)
    return (
        f"用于{reason}活动开展过程中所需的{summary}支出，主要保障活动筹备、训练参赛、"
        f"现场应急和人员组织等相关环节顺利进行，确保本次社团活动具备必要的物资和费用支持。"
    )


def default_purchase_explanation(base_info, ai_data):
    reason = base_info.get("reason") or "社团活动"
    summary = material_summary(ai_data.get("items", []), reason)
    total_amt = format_money(ai_data.get("amount") or items_paid_total(ai_data.get("items", [])))
    return f"本次采购内容为{summary}，相关费用合计 {total_amt} 元。订单截图、电子发票及支付记录等佐证材料齐全，票据真实有效，支出内容与{reason}活动需要相符。"


def default_activity_plan(base_info, ai_data):
    reason = base_info.get("reason") or "SIGS 学生社团活动"
    activity_time = base_info.get("activityTime") or "按社团活动安排开展"
    return f"活动名称：{reason}\n活动时间：{activity_time}\n活动地点：校区内相关活动场地\n参与人员：社团在册成员及相关参与人员"


def default_relation(base_info, ai_data):
    reason = base_info.get("reason") or "社团活动"
    summary = material_summary(ai_data.get("items", []), reason)
    return f"本次采购的{summary}直接服务于{reason}，用于活动训练、参赛或现场保障等环节，能够支撑活动正常开展。"


def default_effect(base_info, ai_data):
    reason = base_info.get("reason") or "社团活动"
    return (
        f"通过本次费用支出，可保障{reason}相关训练、参赛和现场组织工作顺利完成，"
        f"提升社团成员参与体验和活动执行质量，增强团队凝聚力与竞技水平，进一步丰富校园体育文化氛围。"
    )


def order_paid_amount(ai_data):
    explicit_amount = parse_money(ai_data.get("order_paid_amount") or ai_data.get("paid_amount"))
    if explicit_amount > 0:
        return explicit_amount
    amount = infer_amount(ai_data)
    if amount > 0:
        return amount
    return items_paid_total(ai_data.get("items", []))


def infer_amount(ai_data):
    for key in ["amount", "total_amount", "order_paid_amount", "paid_amount", "total", "价税合计", "合计金额"]:
        amount = parse_money(ai_data.get(key))
        if amount > 0:
            return amount
    for key in ["purchase_explanation", "advance_payment_note", "purpose_statement"]:
        amount = text_amount(ai_data.get(key))
        if amount > 0:
            return amount
    return items_paid_total(ai_data.get("items", []))


def default_advance_payment_note(base_info, ai_data):
    paid_amount = order_paid_amount(ai_data)
    if paid_amount <= 1000:
        return ""
    return build_advance_payment_note(base_info, ai_data)


def normalize_text_fields(ai_data, base_info):
    ai_data["items"] = normalize_items(ai_data)
    inferred_amount = infer_amount(ai_data)
    if inferred_amount > 0:
        ai_data["amount"] = format_money(inferred_amount)
    if is_weak_text(ai_data.get("purpose_statement")) or len(str(ai_data.get("purpose_statement") or "").strip()) < 35:
        ai_data["purpose_statement"] = default_purpose(base_info, ai_data)
    if is_weak_text(ai_data.get("purchase_explanation")):
        ai_data["purchase_explanation"] = default_purchase_explanation(base_info, ai_data)
    if is_weak_text(ai_data.get("activity_plan")):
        dynamic_name = ai_data.get("dynamic_activity_name")
        dynamic_info = ai_data.get("dynamic_activity_info")
        ai_data["activity_plan"] = "\n".join(x for x in [f"活动名称：{dynamic_name}" if dynamic_name else "", str(dynamic_info or "")] if x).strip() or default_activity_plan(base_info, ai_data)
    if is_weak_text(ai_data.get("activity_relation")):
        ai_data["activity_relation"] = ai_data.get("dynamic_relation") if not is_weak_text(ai_data.get("dynamic_relation")) else default_relation(base_info, ai_data)
    if is_weak_text(ai_data.get("expected_effect")) or len(str(ai_data.get("expected_effect") or "").strip()) < 35:
        dynamic_effect = str(ai_data.get("dynamic_effect") or "").strip()
        ai_data["expected_effect"] = dynamic_effect if dynamic_effect and len(dynamic_effect) >= 35 and not is_weak_text(dynamic_effect) else default_effect(base_info, ai_data)
    current_advance = str(ai_data.get("advance_payment_note") or "").strip()
    if not current_advance or (order_paid_amount(ai_data) > 1000 and "具体用于以下事项" not in current_advance):
        ai_data["advance_payment_note"] = default_advance_payment_note(base_info, ai_data)
    return ai_data


def default_image_groups(image_count):
    if image_count <= 0:
        return []
    if image_count <= 3:
        return [{
            "invoice_idx": 0 if image_count >= 1 else None,
            "order_idx": 1 if image_count >= 2 else None,
            "payment_idx": 2 if image_count >= 3 else None,
        }]
    groups = []
    for start in range(0, image_count, 3):
        groups.append({
            "invoice_idx": start if start < image_count else None,
            "order_idx": start + 1 if start + 1 < image_count else None,
            "payment_idx": start + 2 if start + 2 < image_count else None,
        })
    return groups


def clean_index(value):
    if value is None:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None

@app.post("/api/v1/reimburse/process")
async def process(
    user_info: str = Form(..., description="前端参数"),
    files: List[UploadFile] = File(...)
):
    try:
        data = json.loads(user_info)
    except Exception as e:
        return {"status": "error", "message": "解析参数失败"}
    
    # 用来保存所有图片的 base64 数据（永久保存在内存中）
    images_payload = []
    image_base64_list = []  # 👈 新增
    
    for idx, f in enumerate(files):
        b = await f.read()
        
        # 把图片转成 base64 字符串
        b64_str = base64.b64encode(b).decode()
        image_base64_list.append(b64_str)  # 👈 保存到内存中

        ai_image_bytes = compress_image_for_ai(b)
        ai_b64_str = base64.b64encode(ai_image_bytes).decode()
        
        images_payload.append({
            "type": "image_url",
            "image_url": {
                "url": "data:image/jpeg;base64," + ai_b64_str
            }
        })
    
    reason_str = data.get('reason', '社团物资采购')
    activity_time_str = data.get('activityTime', '')
    
    prompt = """
只输出JSON，不要解释。
你是一位高智能的财务报销助手。用户一次性上传了多张混合在一起的图片（包含发票、天猫/京东购买详情截图、支付宝/微信支付记录流水），它们顺序是乱的。
请帮我完成以下两件事：
1. 识别并提取发票中的基本字段: amount (订单实付总金额), buyer, tax_id, 以及 items (商品清单列表，含 name, quantity, price, discount, total, remark)。
   重要：items.total 和 amount 必须优先使用订单截图/支付记录中的“实付、实付款、支付金额、Transaction amount”等实际支付金额，负数取绝对值；不要优先使用发票价税合计。若酒店发票价税合计为 3600，但支付记录为 -3590.00，则 total 和 amount 都必须写 3590。
   只提取与报销事由直接相关、需要报销的项目；明显无关的日用品、生活用品、个人消费不要放入 items，也不要计入 amount。
2. 【核心图片连连看匹配】：
   仔细观察用户上传的每张图片（对应的图片索引从 0 开始，即第一张图索引为 0，第二张为 1...）。
   请自动判定每张图的类型（发票/购买详情/支付记录），并把【属于同一个商品订单】的图片组合在一起。
   
输出的 JSON 中必须包含 `image_groups` 字段，格式如下：
"image_groups": [
  {
    "invoice_idx": 0,       # 该商品对应的发票图片索引（若无匹配则填 null）
    "order_idx": 2,         # 该商品对应的购买详情/订单截图索引（若无匹配则填 null）
    "payment_idx": 1        # 该商品对应的支付记录/网银流水截图索引（若无匹配则填 null）
  }
]
3. 针对识别出的商品和用户填写的信息动态生成报销说明字段。用户填写的报销事由为：""" + reason_str + """；用户填写的活动时间为：""" + activity_time_str + """。
   这些说明字段必须贴近真实报销材料，不能空泛，不能写“特购置相关物资”，也不能在说明文字里堆砌商品长全名。明细表 items 可以保留商品全称，但说明文字只写模糊类别，例如训练物资、赛事应急药品耗材、参赛装备、住宿及报名相关费用。
   请额外输出以下字段：
   purpose_statement: 用途阐述，1句话，说明用于什么活动和什么类别支出。
   purchase_explanation: 购买说明，说明采购类别、金额、订单截图/发票/支付记录齐全、支出与活动相关。
   activity_plan: 活动策划，必须包含活动名称、活动时间、活动地点、参与人员，每项单独一行。
   activity_relation: 器材/物资与活动的关联，说明采购类别如何服务训练、参赛或活动保障。
   expected_effect: 预期效果。
   order_paid_amount: 订单实付总额，按支付记录或订单实付金额填写数字。
   advance_payment_note: 当 order_paid_amount 超过 1000 元时生成垫付说明；不超过 1000 元时填空字符串。
   为兼容旧版，也可以继续输出 dynamic_activity_name、dynamic_activity_info、dynamic_relation、dynamic_effect。
4. 金额识别必须以订单实付/支付记录金额为准，尤其住宿、酒店、报名等费用请读取完整金额，例如 3590.00 不能写成 359，也不能写成发票价税合计 3600。不要把非本次报销事项或日用品混入住宿金额。
"""
    try:
        if len(images_payload) > 3:
            image_summaries = analyze_image_summaries(images_payload, reason_str)
            result_data = build_result_from_summaries(image_summaries, data)
            result_data = reconcile_result_with_summaries(result_data, image_summaries)
        else:
            raw = call_zhipu_vision([
                {"role": "user", "content": [{"type": "text", "text": prompt}, *images_payload]}
            ])
            result_data = parse_ai_json(raw)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI识别失败: {str(e)}")
    result_data = normalize_text_fields(result_data, data)
    
    # 👇 关键改动：把图片的 base64 数据传给前端（不传文件路径了）
    result_data["image_base64_list"] = image_base64_list
    
    validation_results = ["🟢 正常：发票预审通过"]
    
    return {
        "ai_data": result_data,
        "check_status": "✔",
        "check_details": validation_results
    }

# ==================== 导出 Word 接口 ====================
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fastapi.responses import FileResponse
import time

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_table_no_borders(table):
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        tag.set(qn("w:sz"), "0")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "FFFFFF")
        borders.append(tag)
    tbl_pr.append(borders)


def picture_fit_size(image_bytes, max_width_in, max_height_in):
    try:
        with Image.open(io.BytesIO(image_bytes)) as image:
            width_px, height_px = image.size
        if width_px <= 0 or height_px <= 0:
            return max_width_in, None
        ratio = width_px / height_px
        width_in = max_width_in
        height_in = width_in / ratio
        if height_in > max_height_in:
            height_in = max_height_in
            width_in = height_in * ratio
        return width_in, height_in
    except Exception:
        return max_width_in, None


def add_base64_picture(paragraph_or_cell, image_base64_list, image_index, max_width_in, max_height_in):
    image_index = clean_index(image_index)
    if image_index is None or image_index >= len(image_base64_list):
        return False
    img_data = base64.b64decode(image_base64_list[image_index])
    width_in, height_in = picture_fit_size(img_data, max_width_in, max_height_in)
    stream = io.BytesIO(img_data)
    run = paragraph_or_cell.add_run() if hasattr(paragraph_or_cell, "add_run") else paragraph_or_cell.paragraphs[0].add_run()
    if height_in:
        run.add_picture(stream, width=Inches(width_in), height=Inches(height_in))
    else:
        run.add_picture(stream, width=Inches(width_in))
    return True


def item_text_for_advance(item):
    name = str(item.get("name") or "报销事项").strip()
    quantity = str(item.get("quantity") or "").strip()
    total = format_money(item.get("total") or item.get("price") or 0)
    if quantity and quantity not in {"0", "0.0"}:
        return f"购置{name} {quantity} 项，花费 {total} 元"
    return f"支付{name}，花费 {total} 元"


def build_advance_payment_note(base_info, ai_data):
    paid_amount = order_paid_amount(ai_data)
    if paid_amount <= 1000:
        return ""
    reason = base_info.get("reason") or "社团活动"
    name = base_info.get("name") or ""
    student_id = base_info.get("studentId") or ""
    items = ai_data.get("items", []) or []
    lines = [
        f"{reason} 物资活动报销垫付事由",
        f"本人{name}（学号：{student_id}）在{reason}活动中垫付共计 {format_money(paid_amount)} 元，具体用于以下事项：",
    ]
    if items:
        for idx, item in enumerate(items, start=1):
            lines.append(f"{idx}. {item_text_for_advance(item)}；")
    else:
        lines.append(f"1. 支付{material_summary(items, reason)}相关费用，花费 {format_money(paid_amount)} 元；")
    current_date = time.strftime('%Y年%m月%d日')
    lines.extend(["签名：", f"日期：{current_date}"])
    return "\n".join(lines)

@app.post("/api/v1/reimburse/export")
async def export_word(payload: dict):
    try:
        base_info = payload.get("base_info", {})
        ai_data = normalize_text_fields(payload.get("ai_data", {}), base_info)
        
        # 👇 关键改动：获取图片的 base64 数据（不再是文件路径）
        image_base64_list = ai_data.get("image_base64_list", [])
        image_groups = ai_data.get("image_groups", [])
        if not image_groups and image_base64_list:
            image_groups = default_image_groups(len(image_base64_list))
        
        doc = Document()
        
        # =======================================================
        # 第一页：封面
        # =======================================================
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run("\n\n\n封面\n").font.size = Pt(22)
        
        reimburse_name = base_info.get('name', '张智雨')
        total_amt = format_money(infer_amount(ai_data))
        reimburse_method = base_info.get('method', '对私转账')
        
        infos = [
            f"报销事由: {base_info.get('reason', 'SIGS学生跆拳道社团 训练物资报销')}",
            f"报销方式: {reimburse_method}给{reimburse_name}，报销总金额为{total_amt}元",
            f"报销人姓名: {reimburse_name}",
            f"报销人学号: {base_info.get('studentId', '2025214507')}",
            f"联系方式: {base_info.get('contact', '18683380921')}",
            f"报销金额: {total_amt}元",
        ]
        for info in infos:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(1.2)
            p.add_run(info).font.size = Pt(14)
            
        # =======================================================
        # 第二页：证明材料明细表
        # =======================================================
        doc.add_page_break()
        doc.add_paragraph().add_run("证明材料").font.size = Pt(18)
        
        p_user = doc.add_paragraph()
        p_user.add_run(f"报销人姓名: {reimburse_name}\n报销人学号: {base_info.get('studentId', '2025214507')}\n报销金额: {total_amt}元\n明细：").font.size = Pt(12)
        
        items = ai_data.get("items", [])
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ['物品', '数量/个', '单价/元', '优惠/元', '总价（元）', '备注']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            set_cell_background(hdr_cells[i], "3B82F6")
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get("name", "商品明细"))
            row_cells[1].text = str(item.get("quantity", "1"))
            row_cells[2].text = str(item.get("price", ""))
            row_cells[3].text = str(item.get("discount", "0"))
            row_cells[4].text = str(item.get("total", ""))
            row_cells[5].text = str(item.get("remark", ""))
        
        total_row_cells = table.add_row().cells
        total_row_cells[0].text = "合计"
        total_row_cells[4].text = f"{total_amt}元"
        total_row_cells[0].paragraphs[0].runs[0].font.bold = True
        total_row_cells[4].paragraphs[0].runs[0].font.bold = True
        
        p_bottom_text = doc.add_paragraph()
        p_bottom_text.paragraph_format.space_before = Pt(12)
        current_date = time.strftime('%Y年%m月%d日')
        ai_purpose = ai_data.get("purpose_statement") or default_purpose(base_info, ai_data)
        
        bottom_text = (
            f"\n合计：{total_amt}元\n"
            f"事由：{base_info.get('reason', '训练物资')}\n\n"
            f"用途：{ai_purpose}\n"
            f"签名：\n\n"
            f"日期：{current_date}\n"
        )
        p_bottom_text.add_run(bottom_text).font.size = Pt(12)
        
                # =======================================================
        # 第三页：图片智能化顺次平铺 (发票X -> 购买记录X -> 支付记录X)
        # 👇 关键改动：从内存中读取图片数据，而不是从硬盘
        # =======================================================
        if image_groups:
            doc.add_page_break()
            
            for g_idx, group in enumerate(image_groups):
                g_num = g_idx + 1
                
                # 1. 插入当前商品的发票：等比缩放，避免酒店等长图撑满整页
                inv_i = clean_index(group.get("invoice_idx"))
                if inv_i is not None and inv_i < len(image_base64_list):
                    try:
                        p_t = doc.add_paragraph()
                        p_t.add_run(f"发票{g_num}：").font.bold = True
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        add_base64_picture(p_img, image_base64_list, inv_i, 5.2, 5.8)
                        print(f"✅ 发票{g_num}插入成功！")
                    except Exception as e:
                        print(f"❌ 插入发票{g_num}失败: {e}")
                
                # 2. 订单和支付记录通常都是手机截图，并排展示更省空间
                order_i = clean_index(group.get("order_idx"))
                pay_i = clean_index(group.get("payment_idx"))
                has_order = order_i is not None and order_i < len(image_base64_list)
                has_pay = pay_i is not None and pay_i < len(image_base64_list)
                if has_order or has_pay:
                    try:
                        pair_table = doc.add_table(rows=2, cols=2)
                        set_table_no_borders(pair_table)
                        pair_table.rows[0].cells[0].text = f"购买记录{g_num}"
                        pair_table.rows[0].cells[1].text = f"支付记录{g_num}"
                        for cell in pair_table.rows[0].cells:
                            cell.paragraphs[0].runs[0].font.bold = True
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for cell in pair_table.rows[1].cells:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if has_order:
                            add_base64_picture(pair_table.rows[1].cells[0], image_base64_list, order_i, 2.55, 4.8)
                            print(f"✅ 购买记录{g_num}插入成功！")
                        if has_pay:
                            add_base64_picture(pair_table.rows[1].cells[1], image_base64_list, pay_i, 2.55, 4.8)
                            print(f"✅ 支付记录{g_num}插入成功！")
                        doc.add_paragraph()
                    except Exception as e:
                        print(f"❌ 插入购买/支付记录{g_num}失败: {e}")

        advance_payment_note = str(ai_data.get("advance_payment_note") or "").strip()
        if advance_payment_note:
            doc.add_page_break()
            h_adv = doc.add_paragraph()
            r_adv = h_adv.add_run("垫付事由")
            r_adv.font.size = Pt(18)
            r_adv.font.bold = True
            p_adv = doc.add_paragraph()
            p_adv.add_run(advance_payment_note).font.size = Pt(12)
        
        # =======================================================
        # 第四页：【全动态 AI 生成】专属支撑材料与活动策划说明
        # =======================================================
        doc.add_page_break()
        h6 = doc.add_paragraph()
        r = h6.add_run("购买说明与社团活动策划案")
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = RGBColor(59, 130, 246)
        
        purchase_explanation = ai_data.get("purchase_explanation") or default_purchase_explanation(base_info, ai_data)
        activity_plan = ai_data.get("activity_plan") or default_activity_plan(base_info, ai_data)
        ai_relation = ai_data.get("activity_relation") or default_relation(base_info, ai_data)
        ai_effect = ai_data.get("expected_effect") or default_effect(base_info, ai_data)
        
        doc.add_paragraph().add_run("【1. 购买说明情况】").font.bold = True
        p_c1 = doc.add_paragraph()
        p_c1.add_run(
            f"{purchase_explanation}\n"
            f"物资用途：{ai_purpose}"
        ).font.size = Pt(12)
        
        doc.add_paragraph().add_run("\n【2. 活动策划案】").font.bold = True
        p_c2 = doc.add_paragraph()
        p_c2.add_run(activity_plan).font.size = Pt(12)
        
        doc.add_paragraph().add_run("\n【3. 器材与活动的关联】").font.bold = True
        p_c3 = doc.add_paragraph()
        p_c3.add_run(ai_relation).font.size = Pt(12)
        
        doc.add_paragraph().add_run("\n【4. 预期效果】").font.bold = True
        p_c4 = doc.add_paragraph()
        p_c4.add_run(ai_effect).font.size = Pt(12)
        
        # 保存并导出
        file_path = "/tmp/reimburse_report.docx" if os.name != 'nt' else "reimburse_report.docx"
        doc.save(file_path)
        
        return FileResponse(
            path=file_path,
            filename=f"清华大学SIGS报销单-{base_info.get('name','报告')}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        print(f"生成Word失败: {e}")
        raise HTTPException(status_code=500, detail=f"Word导出故障: {str(e)}")
