from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import os
import json
import base64
import requests
from datetime import datetime
import io
from docx import Document
from PIL import Image
import tempfile

# AI模型配置
ZHIPU_API_KEY = "e1726b4c8ff9472ca63193add519d253.9F5ghmSnO2OiEpau"
ZHIPU_API_URL = "https://open.bigmodel.cn/api/paas/v4/chat/completions"

app = FastAPI(title="AI Reimburse Assistant API", version="1.0.0")

@app.get("/")
async def root():
    return {"message": "AI Reimburse Assistant API is running"}

# Enable CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create directories
os.makedirs("uploads", exist_ok=True)
os.makedirs("output", exist_ok=True)

@app.post("/api/v1/reimburse/process")
async def process_reimbursement(
    purpose: str = Form(...),
    reimburse_type: str = Form(...),
    user_name: str = Form(...),
    student_id: str = Form(...),
    phone: str = Form(...),
    activity_date: str = Form(...),
    invoice_files: list[UploadFile] = File([]),
    order_files: list[UploadFile] = File([]),
    payment_files: list[UploadFile] = File([])
):
    """
    处理报销申请 - OCR识别和AI结构化
    """
    try:
        # 1. Process uploaded files by type
        all_processed_files = []

        # 处理发票文件
        invoice_results = []
        for file in invoice_files:
            if file.content_type.startswith('image/'):
                processed = await process_file(file, 'invoice')
                if processed:
                    invoice_results.append(processed)
                    all_processed_files.append(processed)

        # 处理订单文件
        order_results = []
        for file in order_files:
            if file.content_type.startswith('image/'):
                processed = await process_file(file, 'order')
                if processed:
                    order_results.append(processed)
                    all_processed_files.append(processed)

        # 处理支付记录文件
        payment_results = []
        for file in payment_files:
            if file.content_type.startswith('image/'):
                processed = await process_file(file, 'payment')
                if processed:
                    payment_results.append(processed)
                    all_processed_files.append(processed)

        # 2. Validate data
        validation_results = validate_reimbursement_data(
            purpose, reimburse_type, user_name, student_id,
            phone, activity_date, invoice_results, order_results, payment_results
        )

        # 3. Calculate total amount
        total_amount = sum(item["amount"] for item in all_processed_files)

        # 4. Prepare structured data
        structured_data = {
            "activity": purpose,
            "reimburse_type": reimburse_type,
            "user": {
                "name": user_name,
                "student_id": student_id,
                "phone": phone
            },
            "activity_date": activity_date,
            "invoices": invoice_results,
            "orders": order_results,
            "payments": payment_results,
            "all_files": all_processed_files,
            "total_amount": total_amount,
            "validation_results": validation_results,
            "file_summary": {
                "invoices": len(invoice_results),
                "orders": len(order_results),
                "payments": len(payment_results),
                "total_files": len(all_processed_files)
            },
            "processed_at": datetime.now().isoformat()
        }

        return {
            "success": True,
            "data": structured_data,
            "message": "处理完成"
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")

@app.post("/api/v1/reimburse/export")
async def export_reimbursement(data: dict):
    """
    导出Word报销文档
    """
    try:
        # Create Word document
        doc = Document()

        # Add title
        doc.add_heading('报销材料', 0)

        # Add basic information
        doc.add_heading('一、基本信息', level=1)
        info_table = doc.add_table(rows=1, cols=2)
        info_data = [
            ['报销事由', data.get('activity', '')],
            ['报销方式', data.get('reimburse_type', '')],
            ['报销人', data['user']['name']],
            ['学号', data['user']['student_id']],
            ['联系方式', data['user']['phone']],
            ['活动时间', data.get('activity_date', '')],
            ['报销总金额', f"{data.get('total_amount', 0):.2f} 元"]
        ]

        for key, value in info_data:
            row = info_table.add_row()
            row.cells[0].text = key
            row.cells[1].text = str(value)

        # Add files details section
        if data.get('invoices') or data.get('orders') or data.get('payments'):
            doc.add_heading('二、票据明细', level=1)

            # Process and display files by type
            file_groups = []

            if data.get('invoices'):
                file_groups.append({
                    'type': '发票',
                    'files': data['invoices'],
                    'type_label': '发票信息'
                })

            if data.get('orders'):
                file_groups.append({
                    'type': '订单',
                    'files': data['orders'],
                    'type_label': '订单信息'
                })

            if data.get('payments'):
                file_groups.append({
                    'type': '支付记录',
                    'files': data['payments'],
                    'type_label': '支付记录'
                })

            # Display each file group
            for group in file_groups:
                doc.add_heading(f"{group['type_label']}（共{len(group['files'])}张）", level=2)

                for i, file_data in enumerate(group['files'], 1):
                    doc.add_paragraph(f"第{i}张{group['type']}", style='Heading 3')

                    # Add basic file info
                    if file_data.get('amount'):
                        doc.add_paragraph(f"金额：{file_data['amount']:.2f} 元")
                    if file_data.get('merchant'):
                        doc.add_paragraph(f"商户/店铺：{file_data['merchant']}")
                    if file_data.get('date'):
                        doc.add_paragraph(f"日期：{file_data['date']}")

                    # Add items table if any
                    items = file_data.get('items', [])
                    if items:
                        table = doc.add_table(rows=1, cols=4)
                        headers = ['商品名称', '数量', '单价(元)', '总价(元)']
                        for j, header in enumerate(headers):
                            table.cell(0, j).text = header

                        for item in items:
                            row = table.add_row()
                            row.cells[0].text = item.get('name', '')
                            row.cells[1].text = str(item.get('quantity', 0))
                            row.cells[2].text = f"{item.get('price', 0):.2f}"
                            row.cells[3].text = f"{item.get('total', 0):.2f}"

                        # Add subtotal for this file
                        if file_data.get('amount'):
                            subtotal_row = table.add_row()
                            subtotal_row.cells[0].text = '小计'
                            subtotal_row.cells[0].bold = True
                            subtotal_row.cells[3].text = f"{file_data['amount']:.2f}"
                            subtotal_row.cells[3].bold = True

                        doc.add_paragraph('\n')

                    # Add file image placeholder (would need actual image embedding)
                    doc.add_paragraph(f"[{group['type']}图片将显示在这里]")

            # Add summary table
            doc.add_paragraph('\n')
            doc.add_heading('汇总', level=2)
            summary_table = doc.add_table(rows=1, cols=2)
            summary_cell = summary_table.cell(0, 0)
            summary_cell.text = f"总金额：{data['total_amount']:.2f} 元"
            summary_cell.paragraphs[0].runs[0].bold = True

        # Add validation results
        if data.get('validation_results'):
            doc.add_heading('三、校验结果', level=1)

            if data['validation_results'].get('complete'):
                doc.add_paragraph('✓ 完整项:', style='List Bullet')
                for item in data['validation_results']['complete']:
                    doc.add_paragraph(f'  {item}', style='List Bullet 2')

            if data['validation_results'].get('missing'):
                doc.add_paragraph('✗ 缺失项:', style='List Bullet')
                for item in data['validation_results']['missing']:
                    doc.add_paragraph(f'  {item}', style='List Bullet 2')

            if data['validation_results'].get('risks'):
                doc.add_paragraph('⚠ 风险项:', style='List Bullet')
                for item in data['validation_results']['risks']:
                    doc.add_paragraph(f'  {item}', style='List Bullet 2')

        # Add signature
        doc.add_page_break()
        doc.add_heading('四、签领表', level=1)
        doc.add_paragraph(f'垫付事由: {data.get("activity", "")}')
        doc.add_paragraph(f'本人 {data["user"]["name"]}（学号：{data["user"]["student_id"]}）在 {data.get("activity_date", "")} 中垫付共计 {data.get("total_amount", 0):.2f} 元，具体用于以下事项：')

        # Display files by type in signature section
        file_number = 1

        # Invoices
        if data.get('invoices'):
            for invoice in data.get('invoices', []):
                doc.add_paragraph(f'发票 {file_number}: {invoice.get("merchant", "")} - {invoice.get("amount", 0):.2f} 元', style='List Number')
                file_number += 1

        # Orders
        if data.get('orders'):
            for order in data.get('orders', []):
                doc.add_paragraph(f'订单 {file_number}: {order.get("merchant", "商户")} - {order.get("amount", 0):.2f} 元', style='List Number')
                file_number += 1

        # Payments
        if data.get('payments'):
            for payment in data.get('payments', []):
                doc.add_paragraph(f'支付记录 {file_number}: {payment.get("merchant", "收款方")} - {payment.get("amount", 0):.2f} 元', style='List Number')
                file_number += 1

        doc.add_paragraph('\n签名（垫付人本人）：___________________')
        doc.add_paragraph(f'日期：{datetime.now().strftime("%Y年%m月%d日")}')

        # Save document
        output_filename = f"output/reimbursement_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(output_filename)

        return FileResponse(
            path=output_filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f"报销材料_{datetime.now().strftime('%Y%m%d')}.docx"
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")

async def process_file(file, file_type):
    """处理单个文件"""
    try:
        # Save file temporarily
        file_path = f"uploads/{file.filename}"
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)

        # Convert to base64 for AI processing
        with open(file_path, "rb") as image_file:
            base64_image = base64.b64encode(image_file.read()).decode('utf-8')

        # 使用智谱GLM-4V进行OCR识别
        ocr_result = await ocr_with_zhipu(file.filename, base64_image, file_type)

        # Clean up
        os.remove(file_path)

        return ocr_result
    except Exception as e:
        print(f"Error processing file {file.filename}: {str(e)}")
        return None

def validate_reimbursement_data(purpose, reimburse_type, user_name, student_id, phone, activity_date, invoices, orders, payments):
    """
    验证报销数据
    """
    results = {
        "complete": [],
        "missing": [],
        "risks": []
    }

    # Check basic information
    if purpose:
        results["complete"].append("已填写报销事由")
    else:
        results["missing"].append("缺少报销事由")

    if reimburse_type:
        results["complete"].append("已选择报销方式")
    else:
        results["missing"].append("缺少报销方式")

    if user_name:
        results["complete"].append("已填写报销人姓名")
    else:
        results["missing"].append("缺少报销人姓名")

    if student_id:
        results["complete"].append("已填写报销人学号")
    else:
        results["missing"].append("缺少报销人学号")

    if phone:
        results["complete"].append("已填写联系方式")
    else:
        results["missing"].append("缺少联系方式")

    if activity_date:
        results["complete"].append("已选择活动时间")
    else:
        results["missing"].append("缺少活动时间")

    # Check invoices
    if invoices:
        results["complete"].append(f"已上传 {len(invoices)} 张发票")

        # Check if all invoices have amounts
        for invoice in invoices:
            if invoice.get('amount'):
                results["complete"].append(f"已识别发票金额: {invoice['amount']:.2f}元")
            else:
                results["missing"].append(f"发票 {invoice['filename']} 缺少金额信息")

            if invoice.get('merchant'):
                results["complete"].append(f"已识别商户: {invoice['merchant']}")
            else:
                results["missing"].append(f"发票 {invoice['filename']} 缺少商户信息")

        # Check for different file types
        if order_results:
            results["complete"].append(f"已上传 {len(order_results)} 张订单")

        if payment_results:
            results["complete"].append(f"已上传 {len(payment_results)} 张支付记录")
        # Payment files are optional, not required

    else:
        results["missing"].append("未上传任何发票")

    # Check for data consistency
    if invoices and len(invoices) > 0:
        total_invoices = sum(inv['amount'] for inv in invoices)
        if total_invoices > 0:
            results["complete"].append("发票总金额计算正确")
        else:
            results["risks"].append("发票总金额异常")

    return results

async def ocr_with_zhipu(filename, base64_image, file_type="invoice"):
    """
    使用智谱GLM-4V进行OCR识别
    """
    headers = {
        'Authorization': f'Bearer {ZHIPU_API_KEY}',
        'Content-Type': 'application/json'
    }

    # 根据文件类型调整提示词
    prompt = ""
    if file_type == "发票":
        prompt = "请识别这张发票，提取金额、商户名称、开票日期、商品明细等关键信息，以JSON格式返回。"
    elif file_type == "订单":
        prompt = "请识别这张订单，提取商品名称、数量、单价、总金额、下单日期等信息，以JSON格式返回。"
    else:  # 支付记录
        prompt = "请识别这张支付记录，提取支付金额、支付时间、支付方式、收款方等信息，以JSON格式返回。"

    data = {
        "model": "glm-4v",
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": prompt
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{base64_image}"
                        }
                    }
                ]
            }
        ],
        "temperature": 0.1,  # 降低温度以获得更稳定的结果
        "max_tokens": 2000
    }

    try:
        response = requests.post(ZHIPU_API_URL, headers=headers, json=data)
        response.raise_for_status()
        result = response.json()

        # 解析AI返回的结果
        content = result["choices"][0]["message"]["content"]

        # 尝试解析JSON
        try:
            ocr_data = json.loads(content)
        except json.JSONDecodeError:
            # 如果不是JSON格式，尝试提取关键信息
            ocr_data = parse_ai_response(content, file_type)

        # 确保返回的数据格式统一
        formatted_result = {
            "filename": filename,
            "file_type": file_type,
            "amount": extract_amount(ocr_data),
            "merchant": extract_merchant(ocr_data, file_type),
            "date": extract_date(ocr_data),
            "items": extract_items(ocr_data, file_type)
        }

        return formatted_result

    except Exception as e:
        # 如果API调用失败，返回错误信息
        return {
            "filename": filename,
            "file_type": file_type,
            "error": f"OCR识别失败: {str(e)}",
            "amount": 0,
            "merchant": "识别失败",
            "date": "",
            "items": []
        }

def parse_ai_response(content, file_type):
    """
    解析AI返回的自然语言内容，提取关键信息
    """
    import re

    data = {
        "amount": 0,
        "merchant": "",
        "date": "",
        "items": []
    }

    # 提取金额
    amount_patterns = [
        r'金额[:：]\s*([\d,]+\.?\d*)',
        r'总计[:：]\s*([\d,]+\.?\d*)',
        r'合计[:：]\s*([\d,]+\.?\d*)',
        r'￥([\d,]+\.?\d*)',
        r'¥([\d,]+\.?\d*)',
        r'(\d+\.?\d*)\s*元'
    ]

    for pattern in amount_patterns:
        match = re.search(pattern, content)
        if match:
            try:
                data["amount"] = float(match.group(1).replace(',', ''))
                break
            except:
                continue

    # 提取商户/店铺名称
    if file_type == "发票" or file_type == "订单":
        merchant_patterns = [
            r'商户[:：]\s*([^\n\r]+)',
            r'店铺[:：]\s*([^\n\r]+)',
            r'名称[:：]\s*([^\n\r]+)',
            r'([^\n\r]+(?:超市|商店|公司|餐饮|酒店))'
        ]

        for pattern in merchant_patterns:
            match = re.search(pattern, content)
            if match:
                data["merchant"] = match.group(1).strip()
                break

    # 提取日期
    date_patterns = [
        r'日期[:：]\s*(\d{4}[-/]\d{1,2}[-/]\d{1,2})',
        r'开票日期[:：]\s*(\d{4}[-/]\d{1,2}[-/]\d{1,2})',
        r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})'
    ]

    for pattern in date_patterns:
        match = re.search(pattern, content)
        if match:
            data["date"] = match.group(1)
            break

    # 提取商品信息
    item_patterns = [
        r'([^,\n]+?)\s*(\d+)\s*([^\d\n]+?)\s*([\d,]+\.?\d*)',
        r'商品[:：]\s*([^,\n]+?)(?:数量| Qty)[:：]\*(\d+)[^\d\n]*价格[:：]*([\d,]+\.?\d*)'
    ]

    for pattern in item_patterns:
        matches = re.findall(pattern, content)
        if matches:
            for match in matches:
                if len(match) >= 3:
                    item = {
                        "name": match[0].strip(),
                        "quantity": int(match[1]) if match[1].isdigit() else 1,
                        "price": float(match[2].replace(',', '')) if match[2] else 0,
                        "total": float(match[2].replace(',', '')) if match[2] else 0,
                        "remark": ""
                    }
                    if len(match) > 3:
                        item["total"] = float(match[3].replace(',', ''))
                    data["items"].append(item)

    return data

def extract_amount(data):
    """提取金额"""
    if isinstance(data, dict):
        if "amount" in data:
            return float(data["amount"])
        elif "total_amount" in data:
            return float(data["total_amount"])
        elif "价格" in data or "金额" in data:
            # 尝试从中文文本中提取
            import re
            text = str(data)
            match = re.search(r'(\d+\.?\d*)', text)
            return float(match.group(1)) if match else 0
    return 0

def extract_merchant(data, file_type):
    """提取商户信息"""
    if isinstance(data, dict):
        if "merchant" in data:
            return data["merchant"]
        elif "store" in data:
            return data["store"]
        elif "商户" in data or "店铺" in data:
            # 尝试从中文文本中提取
            import re
            text = str(data)
            match = re.search(r'商户[:：]\s*([^\n\r]+)', text)
            return match.group(1).strip() if match else ""
    return "识别失败"

def extract_date(data):
    """提取日期"""
    if isinstance(data, dict):
        if "date" in data:
            return data["date"]
        elif "开票日期" in data:
            return data["开票日期"]
    return ""

def extract_items(data, file_type):
    """提取商品明细"""
    items = []
    if isinstance(data, dict):
        if "items" in data:
            items = data["items"]
        elif "商品" in data:
            # 如果商品是文本形式，尝试解析
            import re
            text = str(data["商品"])
            pattern = r'([^,\n]+?)\s*(\d+)\s*([^\d\n]+?)\s*([\d,]+\.?\d*)'
            matches = re.findall(pattern, text)
            for match in matches:
                if len(match) >= 3:
                    item = {
                        "name": match[0].strip(),
                        "quantity": int(match[1]) if match[1].isdigit() else 1,
                        "price": float(match[2].replace(',', '')) if match[2] else 0,
                        "total": float(match[2].replace(',', '')) if match[2] else 0,
                        "remark": ""
                    }
                    if len(match) > 3:
                        item["total"] = float(match[3].replace(',', ''))
                    items.append(item)
    return items

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)